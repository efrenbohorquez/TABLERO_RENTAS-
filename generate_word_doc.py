import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def main():
    doc = Document()
    
    # Título
    title = doc.add_heading('Plan de Aseguramiento de Calidad (QA) - Sistema STAR', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Sección 1
    doc.add_heading('1. Introducción', level=1)
    doc.add_paragraph(
        'Este documento detalla el Sistema de Pruebas de Calidad implementado para el Sistema de '
        'Alerta y Recomendación Territorial (STAR). El objetivo es garantizar la integridad del código, '
        'la fiabilidad de las predicciones y la estabilidad de la interfaz de usuario en futuras actualizaciones.'
    )
    
    # Sección 2
    doc.add_heading('2. Tipos de Pruebas Implementadas', level=1)
    
    doc.add_heading('2.1 Pruebas de Frontend (React / Vite)', level=2)
    doc.add_paragraph('El framework utilizado para las pruebas de interfaz es Vitest en conjunto con React Testing Library.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Ubicación: ').bold = True
    p.add_run('sat-r-dashboard/src/pages/Dashboard.test.tsx')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Cobertura actual: ').bold = True
    p.add_run('Smoke testing del componente principal (Dashboard.tsx). Verifica que la aplicación renderice correctamente los estados de carga y no produzca excepciones inmediatas.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Ejecución: ').bold = True
    p.add_run('npm run test (dentro de la carpeta sat-r-dashboard)')

    doc.add_heading('2.2 Pruebas de Backend y Pipeline (Python)', level=2)
    doc.add_paragraph('Se utiliza pytest para validar las salidas del proceso de Extracción, Transformación, Carga (ETL) y de Machine Learning (XGBoost).')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Ubicación: ').bold = True
    p.add_run('python/tests/test_pipeline.py')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Cobertura actual: ').bold = True
    p.add_run('Verificación de existencia y correcta exportación de archivos JSON esenciales (dashboard_data.json, xgboost_forecast.json, cv_metrics.json).')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Ejecución: ').bold = True
    p.add_run('pytest python/tests')

    # Sección 3
    doc.add_heading('3. Recomendaciones para Futuras Actualizaciones', level=1)
    doc.add_paragraph('Para mantener y extender el control de calidad, se sugiere:', style='Normal')
    
    doc.add_paragraph('1. Pruebas de Componentes (UI): Agregar pruebas para componentes aislados como KpiCard o los selectores de filtro.', style='List Number')
    doc.add_paragraph('2. Pruebas de Regresión (ML): Añadir validaciones en pytest que verifiquen si el MAPE general o el RMSE están dentro de un umbral aceptable antes de desplegar un nuevo modelo en producción.', style='List Number')
    doc.add_paragraph('3. Integración Continua (CI): Integrar estos comandos (npm run test y pytest) en un pipeline de CI/CD (por ejemplo, GitHub Actions) para que se ejecuten automáticamente en cada Pull Request.', style='List Number')
    
    output_path = os.path.join(os.path.dirname(__file__), 'docs', 'PLAN_DE_PRUEBAS.docx')
    doc.save(output_path)
    print(f"Documento Word guardado en: {output_path}")

if __name__ == '__main__':
    main()
